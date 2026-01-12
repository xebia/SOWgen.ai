#!/usr/bin/env python3
"""
Generate a Word document with all UI screenshots of SOWgen.ai application
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def add_heading_with_style(doc, text, level=1):
    """Add a styled heading to the document"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_screenshot(doc, image_path, title, description):
    """Add a screenshot with title and description to the document"""
    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(106, 13, 82)  # Xebia purple color
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add description
    if description:
        desc_para = doc.add_paragraph(description)
        desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add image
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Image not found: {image_path}]")
    
    # Add spacing
    doc.add_paragraph()

def main():
    # Create a new Document
    doc = Document()
    
    # Add title page
    title = doc.add_heading('SOWgen.ai', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('UI Screenshots Documentation', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f'Generated on: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    # Add introduction
    intro_heading = add_heading_with_style(doc, 'Introduction', level=1)
    intro_para = doc.add_paragraph(
        'This document provides a comprehensive visual overview of the SOWgen.ai application. '
        'SOWgen.ai is an enterprise-grade web application designed to streamline how Xebia creates, '
        'manages, and approves Statement of Work (SOW) documents for migration and training services. '
        'The platform combines intelligent automation with professional design to deliver an exceptional '
        'experience for both clients and internal staff.'
    )
    
    doc.add_page_break()
    
    # Add Table of Contents
    toc_heading = add_heading_with_style(doc, 'Table of Contents', level=1)
    doc.add_paragraph('1. Login Page - Client View', style='List Number')
    doc.add_paragraph('2. Login Page - Xebia Team View', style='List Number')
    doc.add_paragraph('3. Services Dashboard', style='List Number')
    doc.add_paragraph('4. SOW Generation Method Selection', style='List Number')
    doc.add_paragraph('5. SOW Form - Basic Information', style='List Number')
    doc.add_paragraph('6. SOW Form - Migration Configuration', style='List Number')
    doc.add_paragraph('7. User Profile', style='List Number')
    doc.add_paragraph('8. Xebia Admin Dashboard', style='List Number')
    
    doc.add_page_break()
    
    # Define screenshots with their descriptions
    screenshots = [
        {
            'path': '/tmp/playwright-logs/01-login-page-client.png',
            'title': '1. Login Page - Client View',
            'description': 'The main login page for client users. Clients can enter their email address and '
                          'organization name to access the platform and create SOWs. The page features a clean, '
                          'modern design with the Xebia branding and highlights three key value propositions: '
                          'Accelerate, Automate, and Scale.'
        },
        {
            'path': '/tmp/playwright-logs/02-login-page-xebia.png',
            'title': '2. Login Page - Xebia Team View',
            'description': 'The login page for Xebia team members (administrators and approvers). Xebia staff '
                          'can sign in using their Xebia email address to access the admin dashboard and manage '
                          'all SOWs across the platform.'
        },
        {
            'path': '/tmp/playwright-logs/03-services-dashboard.png',
            'title': '3. Services Dashboard',
            'description': 'The main dashboard clients see after logging in, displaying all available platform '
                          'services. Users can select from various SCM platforms (GitHub, GitLab, Bitbucket, '
                          'Azure DevOps, etc.) and cloud platforms (AWS, GCP, Azure, Terraform) to generate SOWs. '
                          'The page includes a search bar and category filter for easy navigation, along with '
                          'information about the benefits of automated platform integration.'
        },
        {
            'path': '/tmp/playwright-logs/04-sow-generation-method.png',
            'title': '4. SOW Generation Method Selection',
            'description': 'After selecting a platform (e.g., GitHub), users choose between two SOW generation methods: '
                          'Manual Entry or Automated Import (recommended). The page displays the migration path '
                          'visualization showing the 7-stage journey from source to target platform, including '
                          'Discovery & Analysis, Initial Setup, Repo Migration, CI/CD Migration, CI/CD Implementation, '
                          'Team Training, and Support & Documentation. Integration details are provided at the bottom.'
        },
        {
            'path': '/tmp/playwright-logs/05-sow-form-basic.png',
            'title': '5. SOW Form - Basic Information',
            'description': 'The SOW creation form where users enter basic project information including project name, '
                          'description, and select which services are required (Migration Services and/or Training Services). '
                          'The form features auto-save functionality and allows users to save drafts or submit for approval. '
                          'This view shows the initial project configuration tab.'
        },
        {
            'path': '/tmp/playwright-logs/06-sow-form-migration.png',
            'title': '6. SOW Form - Migration Configuration',
            'description': 'When Migration Services are selected, this section appears showing detailed migration configuration '
                          'options. Users can select the GitHub migration type (Classic, EMU, or GHES), specify the number '
                          'of users to migrate, and view the complete migration path diagram. The form includes repository '
                          'inventory fields (total repositories, public/private counts, size metrics, programming languages) '
                          'and an option to include CI/CD migration. The interactive migration path shows all 7 stages of '
                          'the migration process with visual indicators.'
        },
        {
            'path': '/tmp/playwright-logs/07-user-profile.png',
            'title': '7. User Profile',
            'description': 'The user profile modal displays personal information including name, email, organization, '
                          'user ID, and account type. Users can view their profile information and access the edit '
                          'profile functionality. The profile shows an avatar with user initials and provides a clean '
                          'interface for managing account details.'
        },
        {
            'path': '/tmp/playwright-logs/08-xebia-admin-dashboard.png',
            'title': '8. Xebia Admin Dashboard',
            'description': 'The comprehensive analytics dashboard for Xebia administrators, displaying key metrics '
                          'including Total SOWs, Approved SOWs, Pending Reviews, and Average Approval Time. The dashboard '
                          'features interactive charts showing SOW Status Distribution, Monthly SOW Creation trends, '
                          'SOWs by Client, and a Recent Changes widget. This provides Xebia staff with data-driven '
                          'insights to monitor and optimize the SOW pipeline.'
        }
    ]
    
    # Add each screenshot
    for i, screenshot in enumerate(screenshots, 1):
        add_heading_with_style(doc, screenshot['title'], level=1)
        add_screenshot(doc, screenshot['path'], '', screenshot['description'])
        
        # Add page break after each screenshot except the last one
        if i < len(screenshots):
            doc.add_page_break()
    
    # Add conclusion
    doc.add_page_break()
    conclusion_heading = add_heading_with_style(doc, 'Conclusion', level=1)
    conclusion_para = doc.add_paragraph(
        'This documentation provides a complete visual overview of the SOWgen.ai application, '
        'showcasing its intuitive user interface, comprehensive features, and professional design. '
        'The platform successfully combines intelligent automation with enterprise-grade functionality '
        'to streamline the SOW creation and management process for both clients and Xebia staff.'
    )
    
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run('For more information, visit: https://github.com/xebia/SOWgen.ai')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save the document
    output_path = '/home/runner/work/SOWgen.ai/SOWgen.ai/SOWgen_UI_Screenshots.docx'
    doc.save(output_path)
    print(f'✅ Word document created successfully: {output_path}')
    
    return output_path

if __name__ == '__main__':
    main()
